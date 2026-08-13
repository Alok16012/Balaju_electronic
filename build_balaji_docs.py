from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import csv, random, re

OUT = Path('deliverables')
OUT.mkdir(exist_ok=True)
NAVY = '102A43'; BLUE = '2563EB'; CYAN = '06B6D4'; ORANGE = 'F97316'; INK = '172033'; MUTED = '64748B'; PALE = 'E8F0FE'; LIGHT = 'F5F7FB'; WHITE = 'FFFFFF'; GREEN='15803D'
FONT = 'Aptos'

def rgb(h): return RGBColor.from_string(h)
def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT); run._element.rPr.rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = rgb(color)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=m.find(qn('w:'+tag))
        if x is None: x=OxmlElement('w:'+tag); m.append(x)
        x.set(qn('w:w'), str(val)); x.set(qn('w:type'),'dxa')

def set_cell_text(cell, text, bold=False, color=INK, size=9.2):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
    r=p.add_run(str(text)); font(r,size,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,WHITE,8.8); shade(t.rows[0].cells[i],NAVY)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell_text(cells[i],v,size=8.7); shade(cells[i], WHITE if ridx%2==0 else LIGHT)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def bullet(doc, text, level=0, numbered=False):
    p=doc.add_paragraph(style='List Number' if numbered else 'List Bullet'); p.paragraph_format.left_indent=Inches(.25+.2*level); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.08
    font(p.add_run(text),9.7)
    return p

def para(doc,text,boldlead=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.12
    if boldlead and text.startswith(boldlead):
        font(p.add_run(boldlead),10.2,True,NAVY); font(p.add_run(text[len(boldlead):]),10.2)
    else: font(p.add_run(text),10.2)
    return p

def callout(doc,label,text,fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.columns[0].width=Inches(6.3); c=t.cell(0,0); shade(c,fill); margins(c,140,180,140,180)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); font(p.add_run(label.upper()+'  '),9,True,BLUE); font(p.add_run(text),10,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def setup(doc,title,subtitle,meta):
    sec=doc.sections[0]; sec.page_height=Inches(11); sec.page_width=Inches(8.5); sec.top_margin=sec.bottom_margin=Inches(.78); sec.left_margin=sec.right_margin=Inches(.9); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name=FONT; normal.font.size=Pt(10.2); normal.font.color.rgb=rgb(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12
    for name,size,before,after,color in [('Heading 1',17,16,7,NAVY),('Heading 2',13.5,12,5,BLUE),('Heading 3',11.5,9,4,NAVY)]:
        s=styles[name]; s.font.name=FONT; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=rgb(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for lname in ['List Bullet','List Number']:
        styles[lname].font.name=FONT; styles[lname].font.size=Pt(9.7)
    hdr=sec.header.paragraphs[0]; hdr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(hdr.add_run('BALAJI ELECTRONIC  /  PRODUCT BLUEPRINT'),8,True,MUTED)
    foot=sec.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(foot.add_run('Confidential working document  •  Pune  •  13 August 2026'),8,False,MUTED)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6); font(p.add_run('BALAJI'),11,True,ORANGE)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); font(p.add_run(title),28,True,NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12); font(p.add_run(subtitle),14,False,BLUE)
    for k,v in meta:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(k+': '),9.5,True,MUTED); font(p.add_run(v),9.5)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); bdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'18'); bot.set(qn('w:color'),CYAN); bdr.append(bot); pPr.append(bdr)

def page(doc): doc.add_page_break()
def h1(doc,x): doc.add_heading(x,level=1)
def h2(doc,x): doc.add_heading(x,level=2)
def h3(doc,x): doc.add_heading(x,level=3)

SOURCES=[
('Maharashtra Electronics Corporation — Pune footprint and product range','https://mymec.in/about-us/'),
('Raj Electronics Pune — local categories and delivery/installation cues','https://www.rajelectronicspune.com/'),
('VIANET — Pune B2B distribution, GST billing and dealer pricing','https://www.vianet.co.in/'),
('LeeComp Pune — wholesale computing and component mix','https://www.leecomp.in/'),
('Bajaj Finserv/Croma Pune — mainstream category and store-location overview','https://www.bajajfinserv.in/shop-on-emi-at-croma-store-in-pune'),
('Dribbble reference — Crookeries Ecommerce Website','https://dribbble.com/shots/26163743-Crookeries-Ecommerce-Website'),
]

def build_prd():
    d=Document(); setup(d,'PRODUCT REQUIREMENTS DOCUMENT','A high-trust, animated omnichannel electronics storefront for retail and wholesale buyers',[('Product','Balaji Electronic Commerce Experience'),('Version','1.0 — proof-of-concept definition'),('Primary market','Pune & PCMC, Maharashtra'),('Platforms','Responsive web, mobile web, installable PWA'),('Audience','Founder, product, UX, engineering, merchandising, operations')])
    callout(d,'North star','Help a Pune buyer move from “I need the right electronics product” to a confident order or quote in the fewest meaningful actions—without hiding price, availability, warranty, delivery, or wholesale terms.')
    h2(d,'Executive decision')
    para(d,'Build one storefront with two task modes: Retail and Wholesale. The same inventory powers both, but pricing, quantity controls, trust cues, checkout, documents, and support paths adapt after lightweight role selection. Browsing remains open; sign-in is required only for saved state, protected B2B pricing, quotes, and checkout.')
    h2(d,'Success at a glance')
    table(d,['Outcome','Proof-of-concept target','Guardrail'],[
        ('Findability','80% of test users reach a suitable product in ≤3 interactions','Search and category browse always visible'),
        ('Retail conversion','≥3.0% checkout conversion in usability benchmark','No forced account before cart'),
        ('Wholesale utility','Quote or bulk-cart completion ≥20% of qualified sessions','GST and MOQ visible before commitment'),
        ('Performance','LCP ≤2.5s p75 mobile; INP ≤200ms; CLS ≤0.1','3D never blocks product discovery'),
        ('Trust','Warranty, seller, delivery and returns visible on PDP','No misleading scarcity or fake reviews')],[1.55,2.35,2.4])
    page(d); h1(d,'1. Product vision and scope')
    h2(d,'Problem')
    para(d,'Electronics shopping creates high cognitive load: similar model numbers, technical specifications, uncertain compatibility, hidden installation costs, warranty anxiety, and price comparison. Wholesale buyers add MOQ, tier pricing, GST documentation, repeat ordering, approval, and delivery coordination. Most storefronts force both audiences through the same flow.')
    h2(d,'Vision')
    para(d,'Balaji Electronic becomes Pune’s most helpful digital electronics counter: visually premium, locally credible, unusually clear about money and fulfilment, and fast enough for practical daily use. Motion and 3D communicate product form and system state; they are not decorative delays.')
    h2(d,'In scope for POC')
    for x in ['Responsive landing, category, search, listing, comparison, product detail, cart, checkout, order confirmation and account surfaces.','Retail/Wholesale mode selection and task-based authentication; protected wholesale price tiers and quote workflow.','500 seeded products across 20 categories; dummy prices, stock, media, specifications, ratings and fulfilment.','Bookmarks/wishlist, recently viewed, comparison, saved carts, replenishment lists and shareable inventory QR.','Site Guide Bot that explains navigation, products, compatibility, policies and current screen; safe handoff to human support.','Payment-link-forward checkout, UPI/cards/net banking/wallets/COD configuration, EMI messaging, GST invoice, delivery slot and installation.','Accessible, reduced-motion, low-data and low-power modes.'] : bullet(d,x)
    h2(d,'Explicitly out of scope for POC')
    for x in ['Real payment capture, ERP reconciliation, production courier booking and live tax filing.','Scraping or republishing third-party seller images, prices or claims without permission.','Native iOS/Android apps; PWA covers installability in phase one.','Autonomous bot refunds, price commitments, or warranty adjudication.'] : bullet(d,x)
    h2(d,'Assumptions to validate')
    para(d,'Balaji operates or intends to operate from Pune; final store address, legal entity name, GSTIN, warehouse coverage, service partners, brand authorizations, returns policy and payment provider are not yet supplied. All such fields must remain clearly marked placeholders until verified.')
    page(d); h1(d,'2. Market and audience insight')
    h2(d,'Pune market signal')
    para(d,'The local scan shows broad demand spanning TVs, cooling, major appliances, kitchen products, mobiles, computing and accessories. Pune-based B2B distributors emphasize dealer registration, GST billing, bulk supply, protected pricing and relationship management. Computing specialists add components, custom PCs and peripherals. Community references repeatedly point to Budhwar Peth/Tapkir Galli for components; treat that as directional qualitative evidence, not a verified merchant directory.')
    callout(d,'Product implication','The home page must feel like a premium retailer, while wholesale tools must behave like a procurement console. The catalog taxonomy must accommodate both boxed consumer goods and compatibility-sensitive parts.')
    h2(d,'Primary personas and jobs')
    table(d,['Persona','Job to be done','Decision cues','Fast path'],[
      ('Retail upgrader','Choose a trustworthy appliance or device for home','Final payable price, delivery, installation, warranty, exchange, EMI','Search → compare → PDP → Buy now'),
      ('Value seeker','Find the best fit inside a budget','Deal clarity, price history label, reviews, alternatives','Budget chip → shortlist → compare'),
      ('Retailer/dealer','Replenish multiple SKUs profitably','Dealer price, MOQ, slabs, stock, GST, ETA','Wholesale mode → quick order → quote/pay'),
      ('SME/institution','Buy a configured bundle with approval','Formal quote, tax invoice, delivery plan, support SLA','Use-case bundle → quote → approval'),
      ('Maker/technician','Find compatible parts quickly','Voltage, connector, dimensions, datasheet, pickup','Part search → compatibility → pickup')],[1.25,1.9,1.75,1.65])
    h2(d,'Psychology principles—ethical use')
    for x in ['Recognition over recall: persistent search, recent queries, visually distinct categories and remembered mode.','Choice architecture: lead with use-case and budget filters; progressively disclose expert specifications.','Hick’s law: one dominant CTA per decision state; secondary actions visibly secondary.','Fitts’s law: 44×44 px minimum targets; mobile primary CTA in thumb-reachable sticky bottom bar.','Peak-end effect: make delivery promise, payment confirmation and post-purchase next steps exceptionally clear.','Loss aversion without manipulation: use verified inventory and offer expiry only when sourced from real systems.','Trust transfer: show authorized-brand status only when documented; display warranty owner and installation responsibility.'] : bullet(d,x)
    page(d); h1(d,'3. Experience architecture')
    h2(d,'Global navigation')
    table(d,['Area','Desktop','Mobile'],[
      ('Header','Logo, delivery location, large search, mode switch, account, bookmark, cart','Compact logo, search trigger, mode pill, bookmark, cart'),
      ('Category access','Mega menu grouped by use case and department','Bottom-sheet category browser'),
      ('Primary mobile nav','Not required','Home, Categories, Search, Saved, Account'),
      ('Urgent utility','Running announcement rail with pause control','Single-line swipe rail with pause'),
      ('Help','Floating Site Guide Bot; never covers CTA','Docked help orb above sticky CTA')],[1.25,2.55,2.55])
    h2(d,'Minimum-click task routes')
    for x in ['Known product: focus search on load → autocomplete → product → Buy now → payment.','Exploratory retail: use-case tile → pre-filtered listing → compare up to 4 → product/cart.','Wholesale repeat: sign in → Repeat order → edit quantities → pay or request quote.','Inventory sharing: tap Share inventory → scope selection → QR/link → recipient opens live filtered catalog.','Support: Site Guide suggests the next action based on current page and preserves context during human handoff.'] : bullet(d,x,numbered=True)
    h2(d,'Information architecture')
    para(d,'Home → department/category → search/listing → compare → product → cart → checkout → payment/result → order tracking. Account contains profile, addresses, GST/business profile, orders, quotes, invoices, bookmarks, saved carts, replenishment lists, alerts and support. Content includes buying guides, compatibility explainers, warranty/returns, installation and Pune delivery coverage.')
    h2(d,'Retail vs wholesale behavior')
    table(d,['Capability','Retail mode','Wholesale mode'],[
      ('Entry','Browse immediately; optional sign-in','Browse catalog; verify business to reveal tiers'),
      ('Price','MRP, offer, total savings, EMI','Base/dealer price, GST state, slab tiers, MOQ'),
      ('Cart','Unit quantity, delivery/install add-ons','Case quantity, multi-SKU grid, upload/order by SKU'),
      ('Checkout','Address, slot, payment, invoice','PO/quote, GST, credit terms if approved, multi-address'),
      ('Saved state','Wishlist/bookmarks and alerts','Lists, templates, repeat orders and account pricing')],[1.35,2.45,2.55])
    page(d); h1(d,'4. Landing page specification')
    h2(d,'Narrative sequence')
    table(d,['Zone','Purpose and content','Primary action / behavior'],[
      ('1. Utility rail','Pune delivery promise, store hours, verified offer; horizontally running','Pause; select delivery area'),
      ('2. Hero','3D product constellation around a clear value proposition; Retail/Wholesale task cards','Shop retail / Buy wholesale'),
      ('3. Smart search','Natural-language prompt: “43-inch TV under ₹35k with installation”','Search with voice/text and recent queries'),
      ('4. Category orbit','12 priority categories, 8 secondary; tactile cards with depth','Open pre-filtered category'),
      ('5. Deal theatre','One genuine campaign with price, end time and stock source','View offer; no carousel auto-rotation'),
      ('6. Choose-by-need','New home, office setup, gaming, cooling, creator, repair/maker bundles','Start guided bundle'),
      ('7. Compare strip','Popular comparison trios with plain-language differences','Compare'),
      ('8. Pune trust','Delivery coverage, pickup, installation, GST invoice, support proof','Check PIN code / contact'),
      ('9. Wholesale desk','MOQ, price tiers, quote turnaround, dealer onboarding','Create business account'),
      ('10. Social proof','Verified purchase reviews and project stories','Read proof'),
      ('11. Inventory QR','Live catalog-sharing preview and scope controls','Generate QR'),
      ('12. Footer','Policies, categories, brands, service, business identity','Direct navigation')],[1.05,3.35,2.0])
    h2(d,'Hero interaction rules')
    for x in ['Load a static poster immediately; hydrate the lightweight 3D scene after primary content.','Pointer parallax ≤6 px and object rotation ≤8°; no continuous motion when the user is reading.','Retail and Wholesale cards explain the consequence before selection. Remember mode, but keep switch visible.','Hero CTA remains text-first and high contrast. 3D objects never become the only navigation affordance.','On reduced-motion, replace orbit/scroll choreography with fades ≤150 ms and a static product composition.'] : bullet(d,x)
    h2(d,'Visual direction')
    para(d,'Premium Pune-tech aesthetic: deep navy/ink base, electric blue and cyan for technology, restrained saffron/orange for commerce and local warmth, translucent glass panels over controlled gradients, crisp product photography, generous dark-to-light transitions and rounded geometry. Glass surfaces must retain opaque fallback and at least 4.5:1 text contrast. Use the Dribbble reference for editorial composition, large product imagery and whitespace—not for literal copying.')
    page(d); h1(d,'5. Core journeys and functional requirements')
    h2(d,'Authentication and roles')
    for x in ['Start with Retail / Wholesale intent; allow guest retail browsing and cart.','Retail sign-in: mobile OTP primary; email/password optional; social sign-in only if operationally supported.','Wholesale onboarding: mobile/email OTP, legal/business name, GSTIN where applicable, address, buyer role, expected monthly volume and consent.','States: guest, retail customer, wholesale pending, wholesale verified, wholesale suspended, staff/admin.','Step-up verification for sensitive changes, saved payment instruments, credit terms and high-value orders.','Account merge and role switching must preserve bookmarks and carts without exposing wholesale pricing to unauthorized users.'] : bullet(d,x)
    h2(d,'Search, discovery and comparison')
    for x in ['Typo-tolerant, synonym-aware search across title, brand, model, SKU, category, attributes and compatibility.','Autocomplete groups products, categories, brands, guides and recent queries.','Facets update counts without page reload; selected filters become removable chips; mobile filter is a full-height sheet.','Sort: relevance, price, rating, newest, discount, availability and wholesale margin where authorized.','Compare 2–4 products with sticky difference-only toggle and plain-language “best for” summary.','Zero results offers correction, adjacent category, remove-filter options and human help.'] : bullet(d,x)
    h2(d,'Product detail page')
    for x in ['Above fold: gallery/3D, title/model, rating, final price or protected B2B price, savings, stock, PIN-code ETA, warranty and dominant CTA.','Variant choices show price/stock consequence instantly; unavailable combinations are disabled with reason.','Gallery supports image, short video, 360/3D and dimension overlay; thumbnails remain keyboard accessible.','Compatibility block answers “works with…” and highlights required accessories.','Sticky mobile buy bar: price, Add to cart and Buy now/payment link; bookmark icon near title/gallery.','Below fold: highlights, specs, comparison, offers, delivery/install, reviews, Q&A, documents and alternatives.'] : bullet(d,x)
    h2(d,'Cart, checkout and payment')
    for x in ['Persistent cart across devices after sign-in; clear stock reservation semantics.','Cart shows complete payable estimate, delivery/install, GST state, coupon and savings; no surprise fee at final step.','Express “Buy now” skips cart for a single eligible item.','Checkout uses one page with collapsible completed sections: contact → delivery → invoice → payment.','Payment options: UPI intent/QR, cards, net banking, wallets, configured COD, EMI and a prominent secure payment link.','Payment link is copyable, revocable, time-bound, amount-locked, order-bound and status-synced; warn before sharing.','Failure state preserves cart/order and offers Retry, Change method, or Get help. Never create duplicate orders on retries.'] : bullet(d,x)
    h2(d,'Bookmarks and saved work')
    para(d,'A bookmark is available on every product card and PDP. Guests store locally and are prompted to sync after sign-in. Users can create lists, add notes, share read-only lists, enable price/stock alerts, move items to cart and compare saved items. Wholesale users can convert a list to a repeat-order template or quote request.')
    page(d); h1(d,'6. Site Guide Bot and inventory QR')
    h2(d,'Site Guide Bot')
    para(d,'The bot’s primary role is wayfinding and explanation: “What can I do here?”, “Show wholesale prices”, “How does installation work?”, “Compare these TVs”, or “Where is my order?”. It may summarize approved catalog and policy content, invoke deterministic site actions with confirmation, and hand off context to support.')
    table(d,['Bot capability','Allowed behavior','Safety boundary'],[
      ('Site tour','Contextual guided steps; user-controlled next/skip','Never trap focus or auto-start audio'),
      ('Product help','Grounded summary from structured attributes and approved guides','No invented compatibility, stock or warranty'),
      ('Navigation','Open search, filters, saved items, cart, account page','Confirm before destructive or financial action'),
      ('Order help','Read authorized order status and explain next step','Authenticate; redact sensitive information'),
      ('Commerce','Prepare cart/quote and deep-link to checkout','Cannot place order or promise price autonomously'),
      ('Escalation','Create support ticket/chat handoff with transcript consent','Clear bot identity and easy dismissal')],[1.25,2.6,2.5])
    h3(d,'Bot UX requirements')
    for x in ['Visible “AI site guide” label; never impersonate a human.','Suggested prompts tied to the page; free text always available.','Citations/deep links to product specs and policy sections inside responses.','Conversation can be cleared; transcript retention and consent are explicit.','Marathi, Hindi and English language architecture; POC can launch English first with i18n-ready copy.','Fallback uses deterministic help search when model confidence or retrieval quality is low.'] : bullet(d,x)
    h2(d,'Share entire inventory by QR')
    para(d,'The default QR should encode a short HTTPS URL to a live inventory landing page, not 500 product records. This keeps the QR scannable and inventory current. Authorized users choose scope: full public catalog, category, brand, saved list, quote, or wholesale catalog. Wholesale links require authentication before protected pricing is shown.')
    for x in ['QR preview, downloadable PNG/SVG, copy link, native share and print label.','Link metadata stores creator, scope, filters, expiry, campaign and revocation state.','Public page uses canonical URL, no sensitive query parameters, Open Graph preview and analytics consent.','QR error correction level M by default; minimum printed size tested at intended distance.','Track generated → opened → product viewed → cart/quote, without exposing personal data.'] : bullet(d,x)
    page(d); h1(d,'7. Motion, 3D and micro-interactions')
    h2(d,'Motion system')
    table(d,['Interaction','Motion','Timing / constraint'],[
      ('Button hover','2 px lift, subtle highlight sweep, icon translate','120–180 ms; no layout shift'),
      ('Button press','Scale to 0.98, shadow contracts','80–120 ms; immediate state feedback'),
      ('Bookmark','Icon fill + tiny radial pulse','180–240 ms; announce saved state'),
      ('Add to cart','Product thumbnail arcs toward cart only on capable devices','≤450 ms; instant count update'),
      ('Filter apply','Result skeleton/crossfade; chip settles','150–220 ms; preserve scroll'),
      ('Page transition','Shared element for product image; short fade otherwise','≤300 ms; cancellable'),
      ('Running rail','Constant linear scroll with pause on hover/focus','No critical info exclusively animated'),
      ('3D viewer','Drag rotate, pinch zoom, hotspot focus','User-driven; cap DPR and frame budget')],[1.25,2.55,2.55])
    h2(d,'Performance tiers')
    for x in ['Tier A: static responsive image and CSS micro-interactions for all devices.','Tier B: short product video/360 sprite after consent or interaction.','Tier C: WebGL/glTF 3D only on capable devices and only after primary content.','Stop offscreen animation; respect save-data, reduced-motion, thermal/battery signals where available.','Maintain identical commerce functionality without WebGL, blur, hover or animation.'] : bullet(d,x)
    h2(d,'Accessibility acceptance')
    for x in ['Target WCAG 2.2 AA; keyboard-complete; logical heading and landmark structure.','Visible focus, skip link, semantic controls, form labels, inline errors and error summary.','Touch targets ≥44×44 px; text zoom to 200%; reflow at 320 CSS px without loss.','Contrast: 4.5:1 normal text, 3:1 large text and UI boundaries. Glass does not reduce readability.','Reduced motion disables parallax, auto-orbits, celebratory motion and smooth-scrolling dependence.','Product imagery has meaningful alt text; decorative 3D is hidden from assistive tech; viewer has equivalent specs.'] : bullet(d,x)
    page(d); h1(d,'8. Catalog proof of concept')
    h2(d,'500-SKU composition')
    taxonomy=[('Mobiles',32),('Laptops & tablets',34),('Desktop & components',36),('TV & home theatre',32),('Audio',32),('Refrigerators',24),('Washing machines',24),('Air conditioners & coolers',26),('Kitchen appliances',32),('Small home appliances',28),('Cameras & imaging',16),('Gaming',24),('Networking',24),('CCTV & security',22),('Power & backup',24),('Smart home',20),('Wearables',18),('Mobile accessories',32),('Computer accessories',28),('Electronic components & tools',24)]
    table(d,['Category','Seed SKUs','Hero media'],[(a,b,'3D/360 for top 2–4; image for remainder') for a,b in taxonomy],[2.6,1.15,2.7])
    para(d,'Total: 500 seeded SKUs. The accompanying catalog CSV defines every item, indicative INR price, retail/wholesale eligibility, image direction and a lightweight animation treatment. Names and commercial values are synthetic; brand-like examples must be replaced or licensed before publication.')
    h2(d,'Product data completeness')
    for x in ['Required: SKU, title, slug, category, brand, model, price, tax class, stock, images, short description, key attributes, warranty, fulfilment flags and status.','Wholesale: MOQ, pack size, tier breaks, protected price, lead time and quote eligibility.','Media: 1:1 card image, 4:3 gallery, mobile crop, alt text, focal point, rights owner, license and expiry.','3D: glTF/GLB, Draco/KTX2 where appropriate, poster, dimensions, hotspots and fallback.','Quality gate: no published product without price state, stock state, at least one licensed image, alt text and warranty owner.'] : bullet(d,x)
    h2(d,'Image and animation sourcing')
    para(d,'For the POC, use generated or explicitly licensed images, or manufacturer media with documented permission. Do not hotlink search-result images. Product animation should be authored as reusable treatments—parallax card, hotspot reveal, 360 spin, exploded view, color swap—not 500 heavyweight videos. The seed CSV assigns a treatment per SKU so the UI can demonstrate variety efficiently.')
    page(d); h1(d,'9. Measurement, operations and governance')
    h2(d,'Event taxonomy')
    table(d,['Journey','Core events','Primary metric'],[
      ('Discovery','search_submitted, autocomplete_selected, filter_applied, category_opened','Search success; zero-result rate'),
      ('Evaluation','pdp_viewed, gallery_used, compare_added, bookmark_toggled','PDP→cart; save rate'),
      ('Commerce','cart_added, checkout_started, payment_link_created, payment_result','Conversion; payment success'),
      ('Wholesale','business_signup, verification_state, tier_viewed, quote_requested','Qualified activation; quote rate'),
      ('Help','bot_opened, prompt_used, answer_feedback, handoff_started','Self-serve resolution'),
      ('Share','qr_created, inventory_link_opened, shared_session_converted','Share→commerce rate')],[1.2,3.55,1.65])
    h2(d,'Operational consoles')
    for x in ['Catalog/PIM: products, variants, attributes, media, documents, SEO and publication states.','Inventory/pricing: location stock, safety stock, price lists, offers, tiers and audit history.','Orders/quotes: payment state, fulfilment, invoices, refunds, quote versions and approvals.','Content: landing modules, announcement rail, buying guides, bot knowledge and localization.','Support: tickets, bot handoffs, order context and SLA queue.','Access: least-privilege roles, admin MFA, approval for price/policy changes and immutable audit log.'] : bullet(d,x)
    h2(d,'POC acceptance criteria')
    for x in ['All primary journeys work at 360, 768, 1024 and 1440 px reference widths.','Role switching and authentication states never leak protected pricing.','500 products load from seed data; search, facets, sort, pagination/infinite loading and compare work.','Bookmarking persists for guest and signed-in users; merge behavior is tested.','QR opens the intended inventory scope and honors expiry/revocation.','Bot answers only from approved content, links to source sections and hands off on uncertainty.','Checkout demonstrates final price and payment link clearly; no actual charge in POC.','Automated accessibility, unit, integration and end-to-end checks pass; manual keyboard and screen-reader smoke test passes.','Performance budgets are met on a mid-tier mobile test profile with 3D disabled and enabled separately.'] : bullet(d,x)
    h2(d,'Risks and mitigations')
    table(d,['Risk','Impact','Mitigation'],[
      ('3D and glass overuse','Slow, illegible, distracting','Progressive enhancement, opaque fallback, performance budget'),
      ('Pricing leakage','Commercial and trust harm','Server-side authorization; separate price lists; tests'),
      ('Stale inventory','Failed orders and support burden','Source timestamps, reservations, reconciliation and truthful labels'),
      ('Bot hallucination','Wrong compatibility/policy advice','Retrieval-only knowledge, structured tools, confidence fallback'),
      ('Image rights','Legal/reputation risk','Rights metadata, licensed/generative POC assets, takedown workflow'),
      ('Dark patterns','Short-term lift, long-term distrust','Verified scarcity, clear totals, easy cancellation and consent')],[1.4,2.05,2.95])
    h1(d,'10. Delivery plan and open decisions')
    h2(d,'Phased roadmap')
    table(d,['Phase','Outcome','Indicative scope'],[
      ('0 — Discovery','Validated roles, taxonomy and operating constraints','Stakeholder workshop, content audit, buyer interviews, service map'),
      ('1 — UX prototype','Testable responsive experience','Design system, landing, search, PLP, PDP, role flows, checkout, bot prototype'),
      ('2 — POC build','End-to-end demonstrator with 500 products','Seed catalog, auth sandbox, bookmarks, compare, QR, payment sandbox, analytics'),
      ('3 — Pilot','Controlled Pune launch','Real catalog feeds, payments, fulfilment, support, security and performance hardening'),
      ('4 — Scale','Operational optimization','ERP/PIM integration, personalization, multilingual content, advanced B2B')],[1.2,2.05,3.15])
    h2(d,'Open decisions before production')
    for x in ['Exact Balaji legal/store identity, address, service areas and fulfilment model.','Consumer electronics only vs inclusion of components, repair tools and electrical goods.','Owned stock, marketplace, dropship or hybrid inventory authority.','Wholesale eligibility rules, price tiers, MOQ, credit policy and verification SLA.','Payment provider, COD policy, EMI partners, cancellation/refund policy and payment-link expiry.','Installation partners, warranty ownership, returns categories and reverse-logistics responsibility.','Brand assets, Marathi/Hindi launch requirement, photography rights and final product feed.'] : bullet(d,x)
    h2(d,'Research sources')
    for name,url in SOURCES: para(d,f'{name}: {url}')
    d.save(OUT/'Balaji_Electronic_PRD.docx')

def build_trd():
    d=Document(); setup(d,'TECHNICAL REQUIREMENTS','Architecture, data, security, performance and delivery specification',[('Companion to','Balaji Electronic PRD v1.0'),('Status','POC technical baseline'),('Target','Responsive commerce PWA for Pune retail and wholesale'),('Principle','Fast core commerce; 3D and AI as progressive enhancements')])
    callout(d,'Architecture decision','Use a modular monolith for the POC with strict domain boundaries and event interfaces. It is faster to ship and operate than premature microservices, while preserving clear extraction paths for search, inventory, media and AI.')
    h2(d,'Reference stack')
    table(d,['Layer','Recommended baseline','Reason'],[
      ('Web','Next.js + TypeScript + React; server rendering; PWA','SEO, responsive delivery, shared UI and fast iteration'),
      ('UI','Tailwind/CSS tokens + accessible headless primitives','Consistent system, glass fallbacks, keyboard support'),
      ('Motion/3D','Motion library + Three.js/React Three Fiber; glTF','Controlled micro-motion and progressive 3D'),
      ('API','Typed REST/OpenAPI or tRPC inside modular server','Contract clarity; simple POC deployment'),
      ('Data','PostgreSQL; Redis-compatible cache/queue','Transactional commerce plus fast session/cart/search support'),
      ('Search','Postgres FTS initially; Meilisearch/Typesense extraction path','500-SKU POC simplicity with scalable interface'),
      ('Storage/CDN','S3-compatible object storage + image CDN','Licensed media, transformations and caching'),
      ('Auth','Managed OTP/OIDC provider or audited auth library','Safer identity and step-up verification'),
      ('Payments','India-capable gateway sandbox; signed webhooks','UPI/cards/net banking/links/EMI based on provider'),
      ('AI','Retrieval layer + tool-gated LLM + policy filters','Grounded site help; controlled actions')],[1.2,2.65,2.55])
    page(d); h1(d,'1. System context and domains')
    h2(d,'Logical components')
    for x in ['Web/PWA client: SSR/streamed pages, responsive UI, offline shell, motion and 3D capability detection.','Commerce API: catalog, pricing, inventory, customer, cart, promotion, checkout, orders, payments, quotes, bookmarks and sharing.','Back office: catalog/media, price lists, inventory, content, orders, quotes, user verification and audit.','Integration adapters: payment gateway, SMS/email, address/PIN validation, tax/invoice, courier/installation and future ERP/PIM.','AI gateway: approved knowledge index, structured product retrieval, tool allowlist, safety/PII redaction, evaluation and trace logging.','Observability: structured logs, traces, metrics, frontend vitals, error tracking, audit events and business analytics.'] : bullet(d,x)
    h2(d,'Domain boundaries')
    table(d,['Domain','Owns','Must not trust'],[
      ('Identity & access','User, role, business profile, verification, sessions','Client-provided role or price tier'),
      ('Catalog','Product/variant/category/attribute/media/document','Unlicensed external media'),
      ('Pricing','Price lists, tax display, offers, tier calculation','Cached/client totals at checkout'),
      ('Inventory','Location stock, availability, reservation, lead time','UI availability as fulfillment authority'),
      ('Cart/checkout','Lines, adjustments, address, shipping, final quote','Product price embedded in request'),
      ('Order/payment','Order state, payment intent/link, webhook ledger','Redirect result without signed webhook'),
      ('Sharing','Short link, scope, filters, expiry/revocation','Protected price in public payload'),
      ('AI help','Knowledge, conversations, citations, allowed tool calls','Model-generated commercial facts')],[1.25,2.45,2.7])
    page(d); h1(d,'2. Functional service requirements')
    h2(d,'Identity and authorization')
    for x in ['RBAC roles: guest, retail_customer, wholesale_pending, wholesale_verified, wholesale_suspended, catalog_editor, operations, support, finance_admin, super_admin.','Authorization is enforced server-side per resource and field. Wholesale prices are never serialized for unauthorized clients.','OTP endpoints are rate-limited by IP/device/identity; codes are hashed, short-lived and single-use.','Secure httpOnly, SameSite cookies; CSRF defense; session rotation on authentication and privilege change.','Wholesale approval records actor, timestamp, evidence reference, old/new state and reason.','Admin accounts require MFA; high-risk price, refund and permission actions require re-authentication.'] : bullet(d,x)
    h2(d,'Catalog, search and pricing')
    for x in ['Product supports variants and category-specific attributes; canonical model/SKU uniqueness constraints.','Search index document excludes nonpublic fields; protected wholesale index/price enrichment happens after authorization.','Facets are schema-driven with unit normalization; numeric facets store canonical units.','Price quote endpoint receives SKU/quantity/customer/location/coupon and returns signed calculation ID, line totals, taxes, discounts, expiry and explanations.','Checkout recalculates all totals and rejects stale calculation IDs with a recoverable refresh response.','Availability is location-aware and timestamped; “in stock” semantics are defined per fulfillment mode.'] : bullet(d,x)
    h2(d,'Order and payment state')
    para(d,'Order state: draft → pending_payment → paid/authorized → confirmed → allocated → dispatched/out_for_delivery → delivered, with cancelled/return_requested/returned/refunded branches. Payment state is separate: created → pending → authorized → captured → failed/expired → refunded/partially_refunded. State transitions are explicit, validated and audited.')
    for x in ['Create idempotency key for checkout, order submission, payment creation and refund requests.','Payment links use high-entropy opaque IDs, server-side amount/order binding, configurable expiry, revocation and one-time/limited-use rules.','Webhooks require signature verification, timestamp tolerance, replay protection and durable event ledger.','Redirect success pages poll/read server payment state; they do not mark orders paid.','Store minimal payment metadata; never store raw PAN, CVV or UPI credentials.'] : bullet(d,x)
    page(d); h1(d,'3. Data model')
    h2(d,'Core entities')
    table(d,['Entity','Key fields'],[
      ('User','id, mobile/email, status, locale, created_at'),('BusinessProfile','user_id, legal_name, GSTIN, addresses, verification_state, price_list_id'),('Product','id, slug, title, brand_id, category_id, status, tax_class'),('Variant','id, product_id, sku, model, option_values, dimensions, weight'),('MediaAsset','id, variant_id, type, url, poster_url, alt, rights_owner, license, focal_point'),('PriceList/Price','scope, variant_id, amount, currency, min_qty, starts_at, ends_at'),('Inventory','variant_id, location_id, on_hand, reserved, safety_stock, lead_time'),('Bookmark/List','owner/session, name, visibility; items, note, desired_qty'),('Cart','owner/session, mode, currency, expires_at; lines, calculation_id'),('Quote','business_id, version, lines, totals, expiry, approval_state'),('Order','customer/business, addresses, totals, state, calculation snapshot'),('Payment','order_id, provider_ref, state, amount, method; webhook events'),('ShareLink','token_hash, creator, scope, filter_json, expires_at, revoked_at'),('BotConversation','user/session, consent, messages, citations, tool traces, retention_until')],[1.55,4.85])
    h2(d,'Seed data contract')
    para(d,'The companion CSV is suitable for ingestion after mapping. Production import must validate category, numeric price, stock, slug/SKU uniqueness, image rights, alt text and controlled animation treatment. Dummy images use descriptive placeholders, not third-party copyrighted binaries.')
    h2(d,'Retention and privacy')
    for x in ['Collect only data required for commerce, support, fraud prevention and legal obligations.','Define retention by record type; isolate legal invoice retention from disposable bot/chat telemetry.','Export/delete workflow covers profile, bookmarks and conversations subject to lawful retention.','Analytics uses pseudonymous IDs; advertising/optional analytics waits for consent.','Secrets and sensitive identifiers are encrypted at rest; logs redact OTP, tokens, addresses and payment payloads.'] : bullet(d,x)
    page(d); h1(d,'4. API and integration requirements')
    h2(d,'Representative endpoints')
    table(d,['Method / path','Purpose','Controls'],[
      ('GET /api/catalog/products','Search/browse with facets and cursor','Public fields; cache; abuse limit'),('GET /api/catalog/products/:slug','PDP data and authorized price enrichment','Field-level authorization'),('POST /api/auth/otp/*','Request/verify OTP','Rate limit; hashed code; risk checks'),('POST /api/pricing/quote','Authoritative totals for quantities/context','Auth where needed; signed result'),('POST /api/carts','Create/recover guest or user cart','Signed session; idempotent merge'),('POST /api/checkout/orders','Create order from current calculation','Idempotency; server reprice'),('POST /api/payments/links','Create time-bound payment link','Authorized order owner/operator'),('POST /api/webhooks/payments/:provider','Receive payment event','Signature, replay guard, ledger'),('POST /api/share-links','Create scoped inventory/list URL','Scope policy; expiry; audit'),('POST /api/bot/messages','Grounded answer or approved tool proposal','Auth context; redaction; quotas')],[2.25,2.35,1.8])
    h2(d,'Contract standards')
    for x in ['Versioned OpenAPI contract; generated types; consistent error envelope with user-safe message and trace ID.','Cursor pagination for catalog and operational lists; stable sort keys.','ISO 8601 UTC timestamps; INR stored as integer paise; normalized units with display conversion.','Correlation ID across browser, API, queue and integration; no sensitive data in IDs.','Timeout, retry and circuit-breaker policies per adapter; only retry idempotent operations automatically.'] : bullet(d,x)
    h2(d,'Future integrations')
    para(d,'Define adapters before selecting vendors: InventorySource, PriceSource, PaymentProvider, NotificationProvider, ShippingProvider, InstallationProvider, TaxInvoiceProvider and ProductMediaProvider. The POC ships fake/sandbox implementations behind the same interfaces.')
    page(d); h1(d,'5. Frontend and design-system requirements')
    h2(d,'Responsive baseline')
    for x in ['Mobile-first CSS; content breakpoints driven by layout rather than named devices. Validate at 360, 390, 768, 1024, 1280 and 1440 px.','Server-render catalog/category/PDP content for SEO and fast first paint; hydrate interactive islands selectively.','Shared tokens: color, type, spacing, radius, elevation, blur, motion duration/easing, z-index and container widths.','Components cover navigation, mode switch, search, cards, price, quantity, tier table, badges, filters, compare tray, bookmark, cart, checkout, payment link, bot and QR.','Every component has loading, empty, error, disabled, focus, hover, active and reduced-motion states.'] : bullet(d,x)
    h2(d,'Glass and 3D implementation')
    for x in ['Glass token includes translucent fill, border, shadow and backdrop blur; supports-color/blur fallback is opaque.','Never place low-contrast text directly over changing imagery; add stable scrim or opaque panel.','3D assets: glTF/GLB, compressed geometry/textures, ≤2 MB initial poster+essential payload target, lazy-loaded viewer chunks.','Viewer dynamically caps pixel ratio; pauses render loop when idle/offscreen; disposes GPU resources on unmount.','A static image and complete textual specifications are mandatory fallbacks.'] : bullet(d,x)
    h2(d,'SEO and discoverability')
    for x in ['Unique titles/descriptions, canonical URLs, clean slugs, XML sitemap partitioned by type, robots controls.','Product structured data only when price/availability/review data is truthful and current.','Category content supports Pune intent naturally; no doorway pages or locality keyword stuffing.','Share pages emit safe Open Graph metadata; protected wholesale content is noindex and access-controlled.'] : bullet(d,x)
    page(d); h1(d,'6. Non-functional requirements')
    h2(d,'Performance budgets')
    table(d,['Metric / asset','Target','Test condition'],[
      ('LCP','≤2.5 s p75','Mid-tier mobile, representative 4G, production build'),('INP','≤200 ms p75','Real-user monitoring'),('CLS','≤0.10 p75','Fonts/media dimensions reserved'),('Initial JS','≤200 KB gzip route baseline','Exclude lazy 3D/bot chunks'),('Critical image','Responsive AVIF/WebP; ≤200 KB typical','Correct intrinsic dimensions and priority'),('3D payload','Poster first; viewer lazy; ≤5 MB per featured model goal','Capability-gated and cached'),('API browse p95','≤400 ms from application edge','Warm cache; representative filters'),('Checkout API p95','≤800 ms excluding provider redirect','Includes price/inventory validation')],[1.55,2.15,2.7])
    h2(d,'Availability and resilience')
    for x in ['POC target 99.5%; production SLO to be agreed by business criticality and support coverage.','Graceful degradation: browsing continues if bot/3D/analytics fails; checkout blocks safely if price, inventory or payment authority is unavailable.','Database backup with restore drills; explicit RPO/RTO before pilot.','Queue failed events to dead-letter handling with alerting and replay tooling.','Maintenance mode and feature flags for wholesale prices, payment methods, 3D, bot tools and campaigns.'] : bullet(d,x)
    h2(d,'Accessibility and localization')
    para(d,'WCAG 2.2 AA is a release requirement. Use semantic HTML, complete keyboard support, accessible names/descriptions, predictable focus, announcements for async cart/save states and screen-reader-equivalent 3D data. Externalize all user-facing strings; support INR, Indian number grouping, Asia/Kolkata display and future English/Marathi/Hindi.')
    page(d); h1(d,'7. Security, compliance and abuse prevention')
    h2(d,'Security baseline')
    for x in ['Threat-model authentication, pricing authorization, payment links, webhook handling, admin tools, file upload, share links and bot tools.','Follow current OWASP ASVS-aligned controls; dependency and container scanning in CI; patch SLA by severity.','CSP with nonces/hashes, strict transport security, secure cookies, output encoding and validated redirects.','Rate-limit search, OTP, login, QR creation, bot and checkout; add bot/fraud challenges only when risk signals justify friction.','Validate uploads by type/signature, scan files, strip metadata where appropriate and serve from isolated origin.','Encrypt secrets in managed secret store; rotate; separate dev/stage/prod accounts and data.','Audit admin reads/changes for price, inventory, role, verification, order and refund.'] : bullet(d,x)
    h2(d,'India-specific review gates')
    para(d,'Before production, obtain qualified legal/payment advice for applicable Indian privacy, consumer protection, e-commerce, GST invoicing, payment aggregation, refund and marketing-consent obligations. The system must support policy versioning, consent evidence, seller/business disclosures, grievance/support details, tax invoice data and transparent total pricing; exact obligations depend on Balaji’s legal and operating model.')
    h2(d,'Bot security')
    for x in ['Treat retrieved content and user text as untrusted; defend against prompt injection and data exfiltration.','Tools use server-side authorization and constrained schemas; model never receives reusable secrets.','Financial/destructive actions require deterministic confirmation UI outside free-form chat.','Evaluate factuality, citation correctness, refusal, privacy leakage, cross-role data exposure and unsafe tool calls before release.'] : bullet(d,x)
    page(d); h1(d,'8. Testing and delivery')
    h2(d,'Test pyramid')
    table(d,['Level','Coverage'],[
      ('Unit','Pricing rules, tax rounding, tier selection, state machines, QR scope, permission predicates'),('Component','All states, keyboard behavior, contrast, responsive layout, reduced motion'),('Integration','Database boundaries, cache invalidation, search indexing, payment webhook, notifications'),('E2E','Guest retail, signed retail, wholesale pending/verified, bookmark merge, quote, payment failure/retry, QR'),('Contract','Provider adapters, OpenAPI compatibility, webhook fixtures'),('Performance','Core Web Vitals, load test search/cart/checkout, 3D memory and long-task budgets'),('Security','SAST/SCA, secret scan, DAST, authorization matrix, abuse/rate-limit and pre-pilot penetration test'),('AI evals','Grounding, citations, multilingual prompts, injection, privacy and tool confirmation')],[1.35,5.05])
    h2(d,'CI/CD gates')
    for x in ['Lint/typecheck/unit → build → component/a11y → integration → E2E smoke → dependency/security scan.','Preview deployment for each change with seeded non-sensitive data.','Database migrations are forward-compatible, reviewed, backed up and rollback/roll-forward tested.','Production deploy uses staged rollout and feature flags; payment/webhook changes require sandbox evidence.','Release dashboard shows errors, vitals, payment success, search zero-results and authorization anomalies.'] : bullet(d,x)
    h2(d,'Environments')
    para(d,'Local → preview → staging → production. Staging uses sandbox providers and synthetic identities. Production data must never be copied into lower environments without approved anonymization. Seed scripts are deterministic and versioned.')
    h2(d,'Definition of done')
    for x in ['Acceptance criteria linked to automated or manual evidence.','No critical/high security findings; no known wholesale price exposure path.','Performance and accessibility budgets pass on agreed devices/browsers.','Runbooks exist for payment outage, inventory mismatch, bot disablement, compromised share link and failed deployment.','Operations can edit catalog/content, verify wholesale accounts, trace an order and revoke a share/payment link without engineering.'] : bullet(d,x)
    page(d); h1(d,'9. Build backlog and decisions')
    h2(d,'Suggested epics')
    table(d,['Epic','Key deliverable','Dependency'],[
      ('Foundation','Repo, CI, environments, tokens, shell, observability','Hosting and identity choice'),('Catalog','Schema, seed import, PLP/PDP, media and admin','Taxonomy approval'),('Discovery','Search, facets, autocomplete, compare','Catalog completeness'),('Identity/B2B','OTP, roles, verification, protected tiers','Wholesale policy'),('Saved work','Bookmarks, lists, merge, alerts','Identity and catalog'),('Commerce','Cart, pricing, checkout, sandbox payment link','Payment/provider decisions'),('Share','Scoped short links and QR management','Public URL/domain'),('Site Guide','Knowledge ingestion, grounded answers, safe tools','Approved policies/content'),('Experience','Landing, motion, 3D featured assets, reduced modes','Brand and media rights'),('Hardening','A11y, performance, security, resilience, runbooks','End-to-end feature completion')],[1.3,3.25,1.85])
    h2(d,'Architecture decisions to record')
    for x in ['ADR-001 hosting/runtime and regions; ADR-002 auth provider and session model; ADR-003 catalog/PIM authority; ADR-004 search engine threshold; ADR-005 pricing authority; ADR-006 inventory reservation semantics; ADR-007 payment provider and link model; ADR-008 media rights/CDN; ADR-009 bot model, retention and tool policy; ADR-010 analytics/consent.'] : bullet(d,x)
    h2(d,'Required owner inputs')
    for x in ['Business/legal details, service PIN codes, operating hours and support contacts.','Product feed, authorized brands, stock locations, GST/tax configuration and media rights.','Retail/wholesale price policy, MOQ/slabs, verification, credit and quote approval.','Delivery, installation, warranty, return, cancellation, refund and grievance policies.','Payment methods/providers, analytics consent posture and launch languages.'] : bullet(d,x)
    d.save(OUT/'Balaji_Electronic_Technical_Requirements.docx')

def build_catalog():
    random.seed(23)
    taxonomy=[('Mobiles',32,['5G Smartphone','Feature Phone','Rugged Smartphone']),('Laptops & Tablets',34,['Everyday Laptop','Business Laptop','Gaming Laptop','Android Tablet']),('Desktop & Components',36,['Desktop PC','Motherboard','Graphics Card','NVMe SSD','DDR5 RAM','Power Supply']),('TV & Home Theatre',32,['4K Smart TV','QLED TV','OLED TV','Soundbar','Home Theatre']),('Audio',32,['TWS Earbuds','Wireless Headphones','Bluetooth Speaker','Party Speaker','Microphone']),('Refrigerators',24,['Single Door Refrigerator','Double Door Refrigerator','Side-by-Side Refrigerator']),('Washing Machines',24,['Top Load Washer','Front Load Washer','Semi Automatic Washer']),('Air Conditioners & Coolers',26,['Inverter Split AC','Window AC','Air Cooler','Ceiling Fan']),('Kitchen Appliances',32,['Mixer Grinder','Air Fryer','Microwave Oven','Induction Cooktop','OTG','Electric Kettle']),('Small Home Appliances',28,['Vacuum Cleaner','Steam Iron','Water Heater','Water Purifier','Room Heater']),('Cameras & Imaging',16,['Mirrorless Camera','Action Camera','Instant Camera','Camera Lens']),('Gaming',24,['Gaming Console','Controller','Gaming Monitor','Gaming Keyboard','Gaming Mouse']),('Networking',24,['Wi-Fi 6 Router','Mesh Wi-Fi Kit','Gigabit Switch','USB Wi-Fi Adapter']),('CCTV & Security',22,['IP Camera','DVR Kit','Video Doorbell','Smart Lock']),('Power & Backup',24,['UPS','Home Inverter','Surge Protector','Power Station','Battery']),('Smart Home',20,['Smart Bulb','Smart Plug','Smart Display','IR Remote Hub']),('Wearables',18,['Smartwatch','Fitness Band','GPS Watch']),('Mobile Accessories',32,['Power Bank','USB-C Cable','Fast Charger','Phone Case','Wireless Charger']),('Computer Accessories',28,['Wireless Mouse','Mechanical Keyboard','Webcam','USB-C Hub','Laptop Stand']),('Electronic Components & Tools',24,['Arduino-compatible Board','Sensor Kit','Soldering Station','Digital Multimeter','Breadboard Kit','DC Motor'])]
    brands=['Balaji Select','Astra','VoltEdge','NovaTek','PuneWorks','Orbis','Zenith','Cobalt','PixelArc','IndusTech']
    anims=['card-tilt','slow-360-spin','hotspot-reveal','parallax-depth','color-swap','exploded-view','light-sweep','none-reduced-motion']
    rows=[]; idx=1
    for cat,count,types in taxonomy:
        for j in range(count):
            typ=types[j%len(types)]; brand=brands[(idx*3+j)%len(brands)]; model=f'{re.sub("[^A-Z]","",typ.upper())[:3]}-{1000+idx}'
            base={'Mobiles':6999,'Laptops & Tablets':18999,'Desktop & Components':1499,'TV & Home Theatre':4999,'Audio':799,'Refrigerators':12999,'Washing Machines':8999,'Air Conditioners & Coolers':3999,'Kitchen Appliances':699,'Small Home Appliances':799,'Cameras & Imaging':4999,'Gaming':1299,'Networking':699,'CCTV & Security':899,'Power & Backup':499,'Smart Home':399,'Wearables':999,'Mobile Accessories':199,'Computer Accessories':299,'Electronic Components & Tools':99}[cat]
            price=base + (j%8)*750 + (idx%5)*200; wholesale=round(price*0.82/10)*10
            rows.append({'sku':f'BE-{idx:04d}','title':f'{brand} {typ} {model}','category':cat,'brand':brand,'model':model,'retail_price_inr':price,'wholesale_price_inr':wholesale,'moq':1 if cat not in ['Mobile Accessories','Computer Accessories','Electronic Components & Tools'] else [5,10,20][idx%3],'stock':(idx*7)%83+3,'rating':round(3.8+(idx%12)/10,1),'image_prompt':f'premium studio product photograph of a {typ.lower()}, clean neutral background, ecommerce catalog, front three-quarter view, no logo, 1:1','image_alt':f'{brand} {typ}, front three-quarter product view','animation_treatment':anims[idx%len(anims)],'poc_only':'YES','source_note':'Synthetic proof-of-concept record; replace price, stock, brand and media before production'})
            idx+=1
    with open(OUT/'Balaji_Electronic_500_Product_Seed_Catalog.csv','w',newline='',encoding='utf-8-sig') as f:
        w=csv.DictWriter(f,fieldnames=rows[0].keys()); w.writeheader(); w.writerows(rows)

if __name__=='__main__':
    build_prd(); build_trd(); build_catalog()
